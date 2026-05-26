from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

root = Path(r"C:\Users\sebas\iCloudDrive\Documents\Sebastian\08_Python\00_EAE\TFM-Scraping\Licitaciones")
md_path = root / "MEMORIA_TFM_GOVAI.md"
docx_path = root / "MEMORIA_TFM_GOVAI.docx"

text = md_path.read_text(encoding="utf-8")
lines = text.splitlines()

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Cover page
cover_title = "Memoria Técnica Profesional\nPlataforma GovTech de Inteligencia para Contratación Pública"
p = doc.add_paragraph()
r = p.add_run(cover_title)
r.bold = True
r.font.size = Pt(20)
p.style = doc.styles['Title']

doc.add_paragraph("\nAutor: Sebastián Acosta (editar si aplica)")
doc.add_paragraph("Máster Big Data & Analytics")
doc.add_paragraph("Institución: (Editar universidad/escuela)")
doc.add_paragraph("Fecha: 25 de mayo de 2026")
doc.add_page_break()

# TOC title
h = doc.add_paragraph("Índice")
h.runs[0].bold = True
h.runs[0].font.size = Pt(14)

# Add TOC field
p = doc.add_paragraph()
r = p.add_run()
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'begin')
r._r.append(fldChar)

instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
r._r.append(instrText)

fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'separate')
r._r.append(fldChar)

r = p.add_run("Actualiza este índice en Word: clic derecho > Actualizar campo")
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'end')
r._r.append(fldChar)

doc.add_page_break()

in_code = False
for line in lines:
    s = line.rstrip()

    if s.startswith("```"):
        in_code = not in_code
        continue

    if in_code:
        p = doc.add_paragraph(s)
        p.style = doc.styles['No Spacing']
        continue

    if not s.strip():
        doc.add_paragraph("")
        continue

    if s.startswith("# "):
        doc.add_heading(s[2:].strip(), level=1)
        continue
    if s.startswith("## "):
        doc.add_heading(s[3:].strip(), level=2)
        continue
    if s.startswith("### "):
        doc.add_heading(s[4:].strip(), level=3)
        continue

    if s.startswith("- "):
        doc.add_paragraph(s[2:].strip(), style='List Bullet')
        continue

    # simple markdown table passthrough as text lines
    if s.startswith("|") and s.endswith("|"):
        doc.add_paragraph(s)
        continue

    doc.add_paragraph(s)

doc.save(docx_path)
print(docx_path)
